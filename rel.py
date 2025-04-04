from docx import Document
from docx.shared import Pt


# Função para criar o relatório
def criar_relatorio():
    # Cria um novo documento
    doc = Document()

    # Adiciona o título
    doc.add_heading('Relatório: Modelo de Rede Bayesian com pgmpy', level=1)

    # Adiciona o objetivo
    doc.add_heading('Objetivo', level=2)
    doc.add_paragraph(
        'Este relatório visa descrever o funcionamento, a lógica e os resultados de um código Python que '
        'utiliza a biblioteca pgmpy para criar e realizar inferências em uma Rede Bayesian discreta.'
    )

    # Adiciona o funcionamento do código
    doc.add_heading('Funcionamento do Código', level=2)
    doc.add_paragraph('O código começa importando as bibliotecas necessárias do pgmpy:')

    # Adiciona o trecho de código de importação
    doc.add_paragraph(
        'from pgmpy.models import DiscreteBayesianNetwork\n'
        'from pgmpy.factors.discrete import TabularCPD\n'
        'from pgmpy.inference import VariableElimination',
        style='Code'
    )

    # Adiciona a seção de criação do modelo
    doc.add_heading('1. Criação do Modelo de Rede Bayesian', level=3)
    doc.add_paragraph(
        'O modelo é criado especificando as dependências entre as variáveis:'
    )
    doc.add_paragraph(
        'modelo = DiscreteBayesianNetwork([\'Chuva\', \'Umidade\'], [\'Irrigacao\', \'Umidade\'])',
        style='Code'
    )
    doc.add_paragraph(
        'Isso cria uma rede com duas arestas direcionadas, indicando que "Chuva" e "Irrigação" influenciam "Umidade".'
    )

    # Adiciona a seção de definição das CPDs
    doc.add_heading('2. Definição das CPDs (Distribuições de Probabilidade Condicional)', level=3)
    doc.add_paragraph(
        'As CPDs são definidas para cada variável na rede:'
    )
    doc.add_paragraph(
        'cpd_chuva = TabularCPD(variable=\'Chuva\', variable_card=2, values=[[0.7], [0.3]])\n'
        'cpd_irrigacao = TabularCPD(variable=\'Irrigacao\', variable_card=2, values=[[0.6], [0.4]])\n'
        'cpd_umidade = TabularCPD(\n'
        '    variable=\'Umidade\', variable_card=2,\n'
        '    values=[\n'
        '    [0.8, 0.4, 0.2, 0.1],  # P(Umidade=0)\n'
        '    [0.2, 0.6, 0.8, 0.9]   # P(Umidade=1)\n'
        ' ],\n'
        ' evidence=[\'Chuva\', \'Irrigacao\'],\n'
        ' evidence_card=[2, 2]\n'
        ')',
        style='Code'
    )

    # Adiciona a seção de adição das CPDs ao modelo
    doc.add_heading('3. Adição das CPDs ao Modelo', level=3)
    doc.add_paragraph(
        'As CPDs definidas são adicionadas ao modelo:'
    )
    doc.add_paragraph(
        'modelo.add_cpds(cpd_chuva, cpd_irrigacao, cpd_umidade)',
        style='Code'
    )

    # Adiciona a seção de verificação do modelo
    doc.add_heading('4. Verificação do Modelo', level=3)
    doc.add_paragraph(
        'O modelo é verificado para garantir que a estrutura e as CPDs são consistentes:'
    )
    doc.add_paragraph(
        'modelo.check_model()',
        style='Code'
    )

    # Adiciona a seção de inferência com eliminação de variáveis
    doc.add_heading('5. Inferência com Eliminação de Variáveis', level=3)
    doc.add_paragraph(
        'A inferência é realizada para calcular a probabilidade de "Umidade" dado que "Chuva" ocorreu:'
    )
    doc.add_paragraph(
        'inferencia = VariableElimination(modelo)\n'
        'resultado = inferencia.query(variables=[\'Umidade\'], evidence={\'Chuva\': 1})\n'
        'print(resultado)',
        style='Code'
    )

    # Adiciona a lógica do código
    doc.add_heading('Lógica do Código', level=2)
    doc.add_paragraph(
        'A lógica por trás do código é a seguinte:\n'
        '- A rede bayesiana modela as relações de dependência entre as variáveis "Chuva", "Irrigação" e "Umidade".\n'
        '- As CPDs definem as probabilidades dessas variáveis.\n'
        '- A inferência é usada para calcular a probabilidade de "Umidade" dado um certo estado de "Chuva".'
    )

    # Adiciona os resultados esperados
    doc.add_heading('Resultados Esperados', level=2)
    doc.add_paragraph(
        'O resultado da inferência será a distribuição de probabilidade da variável "Umidade" dado que "Chuva" é igual a 1. '
        'A saída esperada será algo como:'
    )
    doc.add_paragraph(
        '+-----------+------------------+\n'
        '| Umidade   |   phi(Umidade)   |\n'
        '+===========+==================+\n'
        '| Umidade(0)|       0.26       |\n'
        '+-----------+------------------+\n'
        '| Umidade(1)|       0.74       |\n'
        '+-----------+------------------+',
        style='Code'
    )
    doc.add_paragraph(
        'Isso indica que, dado que está chovendo (Chuva=1), a probabilidade de "Umidade" ser 0 é 26%, e a probabilidade de "Umidade" ser 1 é 74%.'
    )

    # Salva o documento
    doc.save('Relatorio_Modelo_Bayesian_ABNT.docx')


# Chama a função para criar o relatório
criar_relatorio()